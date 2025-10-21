# carrera_academica/services/document_service.py
"""
Servicio para generación de documentos dinámicos (Word).
"""
import io
import logging
from typing import Optional, Tuple
from datetime import date

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.utils.text import slugify

from carrera_academica.models import Formulario, PlantillaDocumento, MembreteAnual

logger = logging.getLogger(__name__)


class DocumentService:
    """Servicio para generación de documentos Word dinámicos."""

    @staticmethod
    def generar_documento_dinamico(formulario: Formulario) -> Tuple[Optional[io.BytesIO], Optional[str]]:
        """
        Genera un documento Word personalizado para un formulario.

        Args:
            formulario: Instancia de Formulario

        Returns:
            tuple: (BytesIO con el documento, nombre del archivo) o (None, None)
        """
        ca = formulario.carrera_academica

        plantilla_maestra = PlantillaDocumento.objects.filter(
            tipo_formulario=formulario.tipo_formulario
        ).first()

        membrete = MembreteAnual.objects.filter(
            anio=formulario.anio_correspondiente
        ).first()

        if not plantilla_maestra or not membrete:
            logger.warning(
                f"No se puede generar documento para formulario {formulario.pk}: "
                f"Plantilla={bool(plantilla_maestra)}, Membrete={bool(membrete)}"
            )
            return None, None

        try:
            doc = Document(plantilla_maestra.archivo.path)

            # Reemplazar datos del cuerpo
            contexto = DocumentService._preparar_contexto_reemplazo(
                formulario, ca)
            DocumentService._reemplazar_texto_documento(doc, contexto)

            # Reemplazar datos del encabezado
            DocumentService._reemplazar_encabezado(doc, membrete)

            # Guardar en memoria
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            filename = (
                f"{formulario.tipo_formulario}_{formulario.anio_correspondiente}_"
                f"{slugify(ca.cargo.docente)}.docx"
            )

            logger.info(
                f"Documento generado exitosamente para formulario {formulario.pk}")
            return buffer, filename

        except Exception as e:
            logger.error(
                f"Error generando documento para formulario {formulario.pk}: {e}")
            return None, None

    @staticmethod
    def _preparar_contexto_reemplazo(formulario: Formulario, ca) -> dict:
        """Prepara el diccionario de reemplazos para el documento."""
        return {
            "[DOCENTE_NOMBRE]": str(ca.cargo.docente),
            "[ASIGNATURA]": ca.cargo.asignatura.nombre.title(),
            "[CARGO]": f"{ca.cargo.get_categoria_display()} {ca.cargo.get_caracter_display()}",
            "[ANIO_LECTIVO]": str(formulario.anio_correspondiente),
            "[FECHA_GENERACION]": date.today().strftime("%d/%m/%Y"),
            "[DEDICACION]": ca.cargo.get_dedicacion_display(),
            "[COMISIONES]": "....................",
        }

    @staticmethod
    def _reemplazar_texto_documento(doc: Document, replacements: dict):
        """Reemplaza texto en párrafos y tablas del documento."""
        # Párrafos del cuerpo
        for p in doc.paragraphs:
            for run in p.runs:
                for old_text, new_text in replacements.items():
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

        # Tablas del cuerpo
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            for old_text, new_text in replacements.items():
                                if old_text in run.text:
                                    run.text = run.text.replace(
                                        old_text, new_text)

    @staticmethod
    def _reemplazar_encabezado(doc: Document, membrete: MembreteAnual):
        """Reemplaza el encabezado del documento con logo y frase."""
        header = doc.sections[0].header

        if not header.tables:
            logger.warning("El documento no tiene tabla en el encabezado")
            return

        table = header.tables[0]

        # Celda del logo (columna 0)
        cell_logo = table.cell(0, 0)
        for p in cell_logo.paragraphs:
            if "[LOGO_ANUAL]" in p.text:
                p.text = ""
                p.add_run().add_picture(membrete.logo.path, height=Inches(1.5))

        # Celda de la frase (columna 1)
        cell_frase = table.cell(0, 1)
        for p in cell_frase.paragraphs:
            if "[FRASE_ANUAL]" in p.text:
                p.text = p.text.replace("[FRASE_ANUAL]", membrete.frase)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
