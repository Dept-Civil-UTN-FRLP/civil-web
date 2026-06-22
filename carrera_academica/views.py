# carrera_academica/views.py
from django.http import JsonResponse
import io
import logging
import os
from contextlib import redirect_stderr
from datetime import date, datetime, timedelta

from django.contrib import messages
from django.contrib.auth.decorators import login_required, permission_required
from django.core.exceptions import ValidationError
from django.core.mail import EmailMessage
from django.db.models import Count, Max, Q
from django.http import FileResponse, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.template.loader import render_to_string
from django.utils import timezone
from django.utils.text import slugify
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from pypdf import PdfWriter
from weasyprint import HTML

from carrera_academica.services.document_service import DocumentService
from carrera_academica.services.email_service import EmailService
from carrera_academica.services.pdf_service import PDFService
from carrera_academica.services.ca_service import CAService
from config.pagination import paginate_queryset

from .forms import (
    CargoForm,
    CarreraAcademicaForm,
    EvaluacionForm,
    ExpedienteForm,
    JuntaEvaluadoraForm,
    ResolucionForm,
    JuradoForm,
    VeedorGraduadoForm,
    VeedorEstudianteForm,
    JuradoExternoForm,
    UniversidadForm,
)
from .models import (
    Cargo,
    CarreraAcademica,
    Docente,
    Evaluacion,
    Formulario,
    JuntaEvaluadora,
    MembreteAnual,
    PlantillaDocumento,
    Jurado,
    VeedorEstudiante,
    VeedorGraduado,
    Asignatura,
    JuradoExterno,
    Universidad,
    TipoFormulario,
)

logger = logging.getLogger(__name__)


def replace_text_in_doc(doc, replacements):
    """
    Busca y reemplaza texto en párrafos y tablas, conservando el formato.
    `replacements` es un diccionario con {marcador: texto_nuevo}.
    """
    # Recorremos todos los párrafos del cuerpo del documento
    for p in doc.paragraphs:
        for run in p.runs:
            for old_text, new_text in replacements.items():
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

    # Recorremos todas las tablas del cuerpo del documento
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for old_text, new_text in replacements.items():
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)


def _preparar_contexto_detalle(ca):
    """Prepara la lógica de formularios y años para la vista de detalle."""
    current_year = timezone.now().year

    # Formularios únicos y CV
    formularios_sin_anio = ca.formularios.filter(
        anio_correspondiente__isnull=True
    ).order_by("tipo_formulario")

    form_unicos = list(formularios_sin_anio.filter(
        tipo_formulario__in=["CV", "F01", "F02", "F03"]))

    # Años pendientes de evaluación
    start_year = ca.fecha_inicio.year
    todos_los_anios = set(range(start_year, current_year + 1))
    anios_ya_evaluados = {anio for ev in ca.evaluaciones.all()
                          for anio in ev.anios_evaluados}
    anios_pausados = {item["anio"] for item in ca.anios_pausados}
    anios_pendientes = sorted(
        todos_los_anios - anios_ya_evaluados - anios_pausados)

    # Resumen de años con formularios agrupados
    resumen_anios = ca.get_resumen_anios()
    form_anuales = ca.formularios.filter(
        anio_correspondiente__isnull=False
    ).order_by("anio_correspondiente", "tipo_formulario")

    formularios_por_anio = {}
    for form in form_anuales:
        formularios_por_anio.setdefault(
            form.anio_correspondiente, []).append(form)

    for item in resumen_anios:
        item["formularios"] = formularios_por_anio.get(item["anio"], [])

    # Formularios pendientes de notificación
    tipos_a_notificar = ["F02", "F04", "F05"]
    hay_formularios_pendientes = ca.formularios.filter(
        estado="PEN", tipo_formulario__in=tipos_a_notificar
    ).exists()

    descripciones_formularios = {
        tf.codigo: tf.descripcion
        for tf in TipoFormulario.objects.all()
    }

    return {
        "form_unicos": form_unicos,
        "anios_pendientes_evaluacion": anios_pendientes,
        "hay_formularios_pendientes": hay_formularios_pendientes,
        "resumen_anios": resumen_anios,
        "descripciones_formularios": descripciones_formularios,
    }


@login_required
def dashboard_ca_view(request):
    """Dashboard optimizado de Carrera Académica."""
    # Lógica de Filtros y Búsqueda
    search_query = request.GET.get("q", "")
    estado_filter = request.GET.get("estado", "")

    # Lógica de Formularios Debidos
    current_year = timezone.now().year
    q_formularios_debidos = (
        Q(formularios__anio_correspondiente__lt=current_year)
        | Q(
            formularios__anio_correspondiente=current_year,
            formularios__tipo_formulario="F04",
        )
        | Q(formularios__anio_correspondiente__isnull=True)
    )

    # OPTIMIZACIÓN: Usar el manager personalizado
    carreras_qs = CarreraAcademica.objects.with_related_data().annotate(
        total_formularios_debidos=Count("formularios", filter=q_formularios_debidos),
        formularios_entregados=Count(
            "formularios", filter=Q(formularios__estado="ENT") & q_formularios_debidos
        ),
    )

    # Aplicar filtros
    if search_query:
        carreras_qs = carreras_qs.filter(
            Q(cargo__docente__nombre__icontains=search_query)
            | Q(cargo__docente__apellido__icontains=search_query)
        )
    if estado_filter:
        carreras_qs = carreras_qs.filter(estado=estado_filter)

    # Ordenar
    carreras_qs = carreras_qs.order_by("fecha_vencimiento_actual")

    # ✅ PAGINACIÓN: Aplicar paginación
    page_obj, pagination_context = paginate_queryset(carreras_qs, request, page_size=25)

    # OPTIMIZACIÓN: Ordenar sin queries adicionales
    contexto = {
        "carreras": page_obj,
        "search_query": search_query,
        "estado_filter": estado_filter,
        "estado_choices": CarreraAcademica.ESTADO_CHOICES,
        **pagination_context,
    }

    return render(request, "carrera_academica/dashboard_ca.html", contexto)


@login_required
def detalle_ca_view(request, pk):
    ca = get_object_or_404(CarreraAcademica.objects.with_full_detail(), pk=pk)
    contexto = {
        "ca": ca,
        "form_resolucion": ResolucionForm(),
        "expediente_form": ExpedienteForm(instance=ca),
        "motivo_choices": CarreraAcademica.MOTIVO_ARCHIVO_CHOICES,
        "resultado_choices": CarreraAcademica.RESULTADO_CIERRE_CHOICES,
        **_preparar_contexto_detalle(ca),
    }
    return render(request, "carrera_academica/ca_detail.html", contexto)


@login_required
def subir_formulario_view(request, ca_pk, formulario_pk):
    """Endpoint POST exclusivo para subir archivo a un Formulario. Responde JSON."""
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "Método no permitido."}, status=405)

    ca = get_object_or_404(CarreraAcademica, pk=ca_pk)
    formulario = get_object_or_404(ca.formularios, pk=formulario_pk)
    archivo = request.FILES.get("archivo")

    if not archivo:
        return JsonResponse({"ok": False, "error": "No se recibió ningún archivo."}, status=400)

    if formulario.archivo:
        formulario.archivo.delete(save=False)

    formulario.archivo = archivo
    formulario.estado = "ENT"
    formulario.fecha_entrega = timezone.now().date()
    formulario.save()

    from django.urls import reverse
    ver_url = reverse("serve_private_media", kwargs={"path": formulario.archivo.name[8:]})
    borrar_url = reverse("carrera_academica:borrar_archivo_formulario", kwargs={"pk": formulario.pk})

    return JsonResponse({
        "ok": True,
        "estado": formulario.estado,
        "fecha_entrega": formulario.fecha_entrega.strftime("%d/%m/%Y"),
        "tipo": formulario.tipo_formulario,
        "tipo_display": formulario.get_tipo_formulario_display(),
        "ver_url": ver_url,
        "borrar_url": borrar_url,
    })


@login_required
def iniciar_evaluacion_view(request, pk):
    ca = get_object_or_404(CarreraAcademica, pk=pk)

    # Verificar si se puede iniciar evaluación
    puede, razon = ca.puede_iniciar_evaluacion()
    if not puede:
        messages.error(request, f"No se puede iniciar evaluación: {razon}")
        return redirect("carrera_academica:detalle_ca", pk=ca.pk)

    # --- Lógica para determinar años pendientes ---
    start_year = ca.fecha_inicio.year
    end_year = timezone.now().year
    todos_los_anios = set(range(start_year, end_year + 1))

    anios_ya_evaluados = set()
    for ev in ca.evaluaciones.all():
        for anio in ev.anios_evaluados:
            anios_ya_evaluados.add(anio)
    
    anios_pausados = {item['anio'] for item in ca.anios_pausados}

    anios_pendientes = sorted(
        list(todos_los_anios - anios_ya_evaluados - anios_pausados))

    if request.method == "POST":
        form = EvaluacionForm(request.POST)
        form.fields["anios_a_evaluar"].choices = [(y, y) for y in anios_pendientes]

        if form.is_valid():
            try:
                anios_seleccionados = form.cleaned_data["anios_a_evaluar"]

                # Obtener el siguiente número de evaluación
                from django.db.models import Max

                max_eval = ca.evaluaciones.aggregate(max_num=Max("numero_evaluacion"))[
                    "max_num"
                ]
                nuevo_num = (max_eval or 0) + 1

                # Crear la nueva evaluación
                nueva_evaluacion = Evaluacion(
                    carrera_academica=ca,
                    numero_evaluacion=nuevo_num,
                    anios_evaluados=[int(a) for a in anios_seleccionados],
                )

                # Validar antes de guardar
                nueva_evaluacion.full_clean()
                nueva_evaluacion.save()

                # Crear los formularios asociados
                for tipo in ["F08", "F09", "F10", "F11", "F12"]:
                    Formulario.objects.create(
                        carrera_academica=ca,
                        tipo_formulario=tipo,
                        evaluacion=nueva_evaluacion,
                    )

                es_fetch = request.headers.get("X-Requested-With") == "fetch"
                if es_fetch:
                    from django.template.loader import render_to_string
                    descripciones_formularios = {
                        tf.codigo: tf.descripcion for tf in TipoFormulario.objects.all()
                    }
                    html = render_to_string(
                        "carrera_academica/components/ca/_evaluacion_card.html",
                        {"evaluacion": nueva_evaluacion, "ca": ca,
                         "descripciones_formularios": descripciones_formularios},
                        request=request,
                    )
                    return JsonResponse({"ok": True, "html": html})

                messages.success(
                    request,
                    f"Evaluación N°{nuevo_num} creada, cubriendo los años {', '.join(anios_seleccionados)}.",
                )
                return redirect("carrera_academica:detalle_ca", pk=ca.pk)

            except ValidationError as e:
                logger.warning(f"Error de validación al crear evaluación: {e}")
                if request.headers.get("X-Requested-With") == "fetch":
                    return JsonResponse({"ok": False, "error": " ".join(e.messages)}, status=400)
                for error in e.messages:
                    messages.error(request, error)

            except Exception as e:
                logger.error(f"Error inesperado al crear evaluación: {e}")
                if request.headers.get("X-Requested-With") == "fetch":
                    return JsonResponse({"ok": False, "error": "Error al crear la evaluación."}, status=500)
                messages.error(
                    request, "Error al crear la evaluación. Contacte al administrador."
                )
    else:
        form = EvaluacionForm()
        form.fields["anios_a_evaluar"].choices = [(y, y) for y in anios_pendientes]

    return render(
        request, "carrera_academica/evaluacion_create.html", {"form": form, "ca": ca}
    )


@login_required
def registrar_resolucion_view(request, pk):
    ca = get_object_or_404(CarreraAcademica, pk=pk)
    es_fetch = request.headers.get("X-Requested-With") == "fetch"

    if request.method == "POST":
        form = ResolucionForm(request.POST, request.FILES)
        if form.is_valid():
            nueva_resolucion = form.save(commit=False)
            nueva_resolucion.cargo = ca.cargo
            nueva_resolucion.save()

            objeto = form.cleaned_data["objeto"]

            if objeto in ["alta", "redesignacion", "designacion"]:
                ca.resolucion_designacion = nueva_resolucion
                ca.save()
                if es_fetch:
                    from django.urls import reverse
                    return JsonResponse({
                        "ok": True,
                        "campo": "designacion",
                        "fila_id": "fila-resolucion-designacion",
                        "numero": str(nueva_resolucion),
                        "desvincular_url": reverse("carrera_academica:desvincular_resolucion_ca", kwargs={"pk": ca.pk, "campo": "designacion"}),
                        "file_url": reverse("serve_private_media", kwargs={"path": nueva_resolucion.file.name[8:]}) if nueva_resolucion.file else None,
                        "mensaje": "La resolución se ha vinculado como 'Designación'.",
                    })
                messages.info(request, "La resolución se ha vinculado como 'Designación'.")

            elif objeto == "puesta_funcion":
                ca.resolucion_puesta_en_funcion = nueva_resolucion
                ca.save()
                if es_fetch:
                    from django.urls import reverse
                    return JsonResponse({
                        "ok": True,
                        "campo": "puesta_en_funcion",
                        "fila_id": "fila-resolucion-puesta-en-funcion",
                        "numero": str(nueva_resolucion),
                        "desvincular_url": reverse("carrera_academica:desvincular_resolucion_ca", kwargs={"pk": ca.pk, "campo": "puesta_en_funcion"}),
                        "file_url": reverse("serve_private_media", kwargs={"path": nueva_resolucion.file.name[8:]}) if nueva_resolucion.file else None,
                        "mensaje": "La resolución se ha vinculado como 'Puesta en Función'.",
                    })
                messages.info(request, "La resolución se ha vinculado como 'Puesta en Función'.")

            elif objeto == "prorroga_ca":
                exito, mensaje = CAService.aplicar_prorroga(
                    ca,
                    nueva_fecha=form.cleaned_data.get("nueva_fecha_vencimiento"),
                    dias=form.cleaned_data.get("prorroga_dias"),
                )
                if es_fetch:
                    if exito:
                        ca.refresh_from_db()
                        return JsonResponse({
                            "ok": True,
                            "campo": "prorroga",
                            "nueva_fecha": ca.fecha_vencimiento_actual.strftime("%d/%m/%Y"),
                            "mensaje": mensaje,
                        })
                    return JsonResponse({"ok": False, "error": mensaje}, status=400)
                if exito:
                    messages.success(request, mensaje)
                else:
                    messages.error(request, mensaje)
                return redirect("carrera_academica:detalle_ca", pk=ca.pk)

            elif objeto == "licencia_alta":
                ca.estado = "STB"
                ca.save()
                if es_fetch:
                    return JsonResponse({
                        "ok": True,
                        "campo": "estado",
                        "estado_display": ca.get_estado_display(),
                        "mensaje": f"Resolución de '{nueva_resolucion.get_objeto_display()}' registrada.",
                    })

            elif objeto == "licencia_baja":
                ca.estado = "ACT"
                ca.save()
                if es_fetch:
                    return JsonResponse({
                        "ok": True,
                        "campo": "estado",
                        "estado_display": ca.get_estado_display(),
                        "mensaje": f"Resolución de '{nueva_resolucion.get_objeto_display()}' registrada.",
                    })
            else:
                ca.save()

            messages.success(
                request,
                f"Resolución de '{nueva_resolucion.get_objeto_display()}' registrada exitosamente.",
            )
        elif es_fetch:
            errores = {f: e.as_text() for f, e in form.errors.items()}
            return JsonResponse({"ok": False, "error": "Datos inválidos.", "errores": errores}, status=400)

    return redirect("carrera_academica:detalle_ca", pk=ca.pk)


@login_required
def crear_ca_view(request):
    ca_form = CarreraAcademicaForm()
    cargo_form = CargoForm()

    if request.method == "POST":
        if "submit_existente" in request.POST:
            form = CarreraAcademicaForm(request.POST)
            if form.is_valid():
                try:
                    cargo_seleccionado = form.cleaned_data["cargo"]

                    # Crear la CA manualmente con las fechas del cargo
                    nueva_ca = CarreraAcademica(
                        cargo=cargo_seleccionado,
                        numero_expediente=form.cleaned_data["numero_expediente"],
                        fecha_inicio=cargo_seleccionado.fecha_inicio,
                        fecha_vencimiento_original=cargo_seleccionado.fecha_vencimiento,
                    )

                    # Validar antes de guardar
                    nueva_ca.full_clean()
                    nueva_ca.save()

                    messages.success(
                        request,
                        f"Carrera Académica iniciada para el cargo de {cargo_seleccionado}.",
                    )
                    return redirect("carrera_academica:dashboard_ca")

                except ValidationError as e:
                    # Manejar errores de validación
                    logger.warning(f"Error de validación al crear CA: {e}")
                    for field, errors in e.message_dict.items():
                        for error in errors:
                            messages.error(request, f"{field}: {error}")
                    ca_form = form

                except Exception as e:
                    logger.error(f"Error inesperado al crear CA: {e}")
                    messages.error(
                        request,
                        "Error al crear la Carrera Académica. Contacte al administrador.",
                    )
                    ca_form = form
            else:
                ca_form = form

        elif "submit_nuevo" in request.POST:
            form = CargoForm(request.POST)
            if form.is_valid():
                try:
                    # Crear el nuevo cargo
                    nuevo_cargo = form.save()

                    # Crear la CA asociada
                    nueva_ca = CarreraAcademica(
                        cargo=nuevo_cargo,
                        fecha_inicio=nuevo_cargo.fecha_inicio,
                        fecha_vencimiento_original=nuevo_cargo.fecha_vencimiento,
                    )

                    nueva_ca.full_clean()
                    nueva_ca.save()

                    messages.success(
                        request,
                        f"Nuevo cargo y Carrera Académica creados para {nuevo_cargo.docente}.",
                    )
                    return redirect("carrera_academica:dashboard_ca")

                except ValidationError as e:
                    logger.warning(f"Error de validación al crear cargo y CA: {e}")
                    for field, errors in e.message_dict.items():
                        for error in errors:
                            messages.error(request, f"{field}: {error}")
                    cargo_form = form

                except Exception as e:
                    logger.error(f"Error inesperado al crear cargo y CA: {e}")
                    messages.error(
                        request, "Error al crear el cargo. Contacte al administrador."
                    )
                    cargo_form = form
            else:
                cargo_form = form

    contexto = {
        "ca_form": ca_form,
        "cargo_form": cargo_form,
    }
    return render(request, "carrera_academica/ca_create.html", contexto)


@login_required
def editar_junta_view(request, pk):
    """Vista optimizada para editar junta."""
    # ✅ OPTIMIZACIÓN: Precargar relaciones necesarias
    ca = get_object_or_404(
        CarreraAcademica.objects.select_related("cargo__docente", "cargo__asignatura"),
        pk=pk,
    )

    junta, created = JuntaEvaluadora.objects.get_or_create(carrera_academica=ca)

    if request.method == "POST":
        form = JuntaEvaluadoraForm(request.POST, instance=junta)
        if form.is_valid():
            form.save()
            messages.success(
                request, "La Junta Evaluadora ha sido actualizada exitosamente."
            )
            return redirect("carrera_academica:detalle_ca", pk=ca.pk)
    else:
        form = JuntaEvaluadoraForm(instance=junta)

    contexto = {
        "form": form,
        "ca": ca,
        "categoria_choices": Cargo.CATEGORIA_CHOICES,
        "dedicacion_choices": Cargo.DEDICACION_CHOICES,
    }
    return render(request, "carrera_academica/junta_update.html", contexto)


@login_required
def desvincular_resolucion_ca_view(request, pk, campo):
    """Desvincula una resolución (designación o puesta en función) de una CA."""
    if request.method != "POST":
        return redirect("carrera_academica:detalle_ca", pk=pk)

    ca = get_object_or_404(CarreraAcademica, pk=pk)
    es_fetch = request.headers.get("X-Requested-With") == "fetch"

    campos_validos = {
        "designacion": ("resolucion_designacion", "Designación"),
        "puesta_en_funcion": ("resolucion_puesta_en_funcion", "Puesta en Función"),
    }

    if campo not in campos_validos:
        if es_fetch:
            return JsonResponse({"ok": False, "error": "Campo no válido."}, status=400)
        messages.error(request, "Campo no válido.")
        return redirect("carrera_academica:detalle_ca", pk=pk)

    field_name, label = campos_validos[campo]
    setattr(ca, field_name, None)
    ca.save(update_fields=[field_name])
    if es_fetch:
        return JsonResponse({"ok": True, "label": label})
    messages.success(request, f"Resolución de {label} desvinculada correctamente.")
    return redirect("carrera_academica:detalle_ca", pk=pk)


@login_required
def asignar_expediente_view(request, pk):
    ca = get_object_or_404(CarreraAcademica, pk=pk)
    if request.method == "POST":
        form = ExpedienteForm(request.POST, instance=ca)
        if form.is_valid():
            form.save()
            if request.headers.get("X-Requested-With") == "fetch":
                return JsonResponse({"ok": True, "numero_expediente": ca.numero_expediente})
            messages.success(request, "Número de expediente actualizado correctamente.")
        elif request.headers.get("X-Requested-With") == "fetch":
            return JsonResponse({"ok": False, "error": "Número de expediente inválido."}, status=400)
    return redirect("carrera_academica:detalle_ca", pk=ca.pk)


@login_required
@permission_required('carrera_academica.view_carreraacademica', raise_exception=True)
def docentes_filtrados_api_view(request):
    """API optimizada para filtrar docentes."""
    # ✅ OPTIMIZACIÓN: select_related para evitar queries adicionales
    queryset = (
        Docente.objects.filter(cargo_docente__caracter__in=["ord", "reg"])
        .select_related()
        .distinct()
    )

    categoria_seleccionada = request.GET.get("categoria")
    dedicacion_seleccionada = request.GET.get("dedicacion")

    # Lógica de Filtro Jerárquico para CATEGORÍA
    if categoria_seleccionada:
        categorias_orden = ["jtp", "adj", "aso", "tit"]
        try:
            start_index = categorias_orden.index(categoria_seleccionada)
            categorias_validas = categorias_orden[start_index:]
            queryset = queryset.filter(cargo_docente__categoria__in=categorias_validas)
        except ValueError:
            pass

    # Lógica de Filtro Jerárquico para DEDICACIÓN
    if dedicacion_seleccionada:
        dedicaciones_orden = ["ds", "se", "de"]
        try:
            start_index = dedicaciones_orden.index(dedicacion_seleccionada)
            dedicaciones_validas = dedicaciones_orden[start_index:]
            queryset = queryset.filter(
                cargo_docente__dedicacion__in=dedicaciones_validas
            )
        except ValueError:
            pass

    # ✅ OPTIMIZACIÓN: only() para traer solo los campos necesarios
    docentes_list = list(
        queryset.only("id", "apellido", "nombre").values("id", "apellido", "nombre")
    )

    for docente in docentes_list:
        docente["full_name"] = (
            f"{docente['apellido'].upper()}, {docente['nombre'].title()}"
        )

    return JsonResponse({"docentes": docentes_list})


@login_required
def finalizar_ca_view(request, pk):
    ca = get_object_or_404(CarreraAcademica, pk=pk)
    es_fetch = request.headers.get("X-Requested-With") == "fetch"

    if request.method == "POST":
        resultado = request.POST.get("resultado_cierre")
        observaciones = request.POST.get("observaciones_cierre", "")

        if not resultado:
            if es_fetch:
                return JsonResponse({"ok": False, "error": "Debe seleccionar un resultado de cierre."}, status=400)
            messages.error(request, "Debe seleccionar un resultado de cierre")
            return redirect("carrera_academica:finalizar_ca", pk=pk)

        nueva_fecha = None
        if resultado == "aprobada_redesigna":
            fecha_str = request.POST.get("nueva_fecha_vencimiento")
            try:
                nueva_fecha = datetime.strptime(fecha_str, "%Y-%m-%d").date()
            except (ValueError, TypeError):
                if es_fetch:
                    return JsonResponse({"ok": False, "error": "Formato de fecha inválido."}, status=400)
                messages.error(request, "Formato de fecha inválido")
                return redirect("carrera_academica:finalizar_ca", pk=pk)

        exito, mensaje = CAService.finalizar(
            ca, resultado, observaciones, nueva_fecha, request.user)

        if not exito:
            if es_fetch:
                return JsonResponse({"ok": False, "error": mensaje}, status=400)
            messages.error(request, mensaje)
            return redirect("carrera_academica:finalizar_ca", pk=pk)

        # Redesignación: crear nueva CA y redirigir
        if mensaje.startswith("redesignacion:"):
            from django.urls import reverse
            nueva_ca_pk = mensaje.split(":")[1]
            if es_fetch:
                return JsonResponse({
                    "ok": True,
                    "redesignacion": True,
                    "redirect_url": reverse("carrera_academica:detalle_ca", kwargs={"pk": nueva_ca_pk}),
                    "mensaje": f"CA finalizada y redesignada. Nuevo expediente #{nueva_ca_pk} creado.",
                })
            messages.success(request, f"CA finalizada y redesignada. Nuevo expediente #{nueva_ca_pk} creado.")
            return redirect("carrera_academica:detalle_ca", pk=nueva_ca_pk)

        if es_fetch:
            return JsonResponse({
                "ok": True,
                "estado_display": ca.get_estado_display(),
                "mensaje": f"Expediente finalizado: {mensaje}",
            })
        messages.success(request, f"Expediente finalizado: {mensaje}")
        return redirect("carrera_academica:detalle_ca", pk=pk)

    context = {
        "ca": ca,
        "resultado_choices": CarreraAcademica.RESULTADO_CIERRE_CHOICES,
    }
    return render(request, "carrera_academica/finalizar_ca.html", context)


@login_required
def consolidar_pdf_view(request, pk):
    """Vista para consolidar expediente en PDF."""
    ca = get_object_or_404(CarreraAcademica, pk=pk)

    output_buffer, errores = PDFService.consolidar_expediente(ca)

    if not output_buffer:
        messages.error(request, "No se pudo generar el PDF consolidado")
        return redirect("carrera_academica:detalle_ca", pk=ca.pk)

    for error in errores:
        messages.warning(request, error)

    response = HttpResponse(output_buffer, content_type="application/pdf")
    response["Content-Disposition"] = (
        f'attachment; filename="expediente_{slugify(ca.cargo.docente)}.pdf"'
    )
    return response


@login_required
def generar_propuesta_jurado_view(request, pk):
    """Vista para generar PDF de propuesta de jurado."""
    ca = get_object_or_404(CarreraAcademica, pk=pk)

    signature_path = "/static/images/firma_holografica.png"
    pdf_file = PDFService.generar_propuesta_jurado(ca, signature_path)

    if not pdf_file:
        messages.error(request, "No se pudo generar la propuesta de jurado")
        return redirect("carrera_academica:detalle_ca", pk=ca.pk)

    response = HttpResponse(pdf_file, content_type="application/pdf")
    response["Content-Disposition"] = (
        f'attachment; filename="propuesta_jurado_{slugify(ca.cargo.docente)}.pdf"'
    )
    return response


@login_required
def notificar_pendientes_view(request, pk):
    """Vista para notificar formularios pendientes."""
    ca = get_object_or_404(CarreraAcademica, pk=pk)

    exito, mensaje = EmailService.enviar_primera_notificacion(ca)

    if exito:
        messages.success(request, mensaje)
    else:
        messages.error(request, mensaje)

    return redirect("carrera_academica:detalle_ca", pk=ca.pk)


@login_required
def descargar_plantilla_view(request, pk):
    """Vista para descargar plantilla de formulario."""
    formulario = get_object_or_404(Formulario, pk=pk)
    tipos_dinamicos = ["F06", "F07", "F13", "ENC", "F04", "F05"]

    if formulario.tipo_formulario in tipos_dinamicos:
        buffer, filename = DocumentService.generar_documento_dinamico(formulario)

        if buffer:
            return FileResponse(buffer, as_attachment=True, filename=filename)
        else:
            messages.error(
                request,
                "No se pudo generar el documento. Verifique plantillas y membretes.",
            )
            return redirect("carrera_academica:detalle_ca", pk=formulario.carrera_academica.pk)
    else:
        # Lógica para plantillas estáticas
        plantilla = PlantillaDocumento.objects.filter(
            tipo_formulario=formulario.tipo_formulario
        ).first()

        if plantilla and plantilla.archivo:
            return FileResponse(
                plantilla.archivo.open("rb"),
                as_attachment=True,
                filename=plantilla.archivo.name,
            )
        else:
            messages.error(
                request, f"No se encontró plantilla para {formulario.tipo_formulario}."
            )
            return redirect("carrera_academica:detalle_ca", pk=formulario.carrera_academica.pk)


@login_required
def notificar_junta_view(request, pk):
    """Vista para notificar a la junta evaluadora."""
    evaluacion = get_object_or_404(Evaluacion, pk=pk)
    ca = evaluacion.carrera_academica

    emails_enviados, errores = EmailService.enviar_notificacion_junta(evaluacion)

    if emails_enviados > 0:
        messages.success(
            request,
            f"Se han enviado {emails_enviados} correos a los miembros de la junta.",
        )

    for error in errores:
        messages.warning(request, error)

    return redirect("carrera_academica:detalle_ca", pk=ca.pk)


@login_required
def agendar_evaluacion_view(request, pk):
    evaluacion = get_object_or_404(Evaluacion, pk=pk)
    es_fetch = request.headers.get("X-Requested-With") == "fetch"

    if request.method == "POST":
        fecha_str = request.POST.get("fecha_evaluacion")

        if fecha_str:
            evaluacion.fecha_evaluacion = fecha_str
            evaluacion.save()
            if es_fetch:
                return JsonResponse({
                    "ok": True,
                    "fecha_display": evaluacion.fecha_evaluacion.strftime("%d/%m/%Y %H:%M"),
                })
            messages.success(
                request,
                f"Se agendó la fecha para la Evaluación N°{evaluacion.numero_evaluacion}.",
            )
        else:
            evaluacion.fecha_evaluacion = None
            evaluacion.save()
            if es_fetch:
                return JsonResponse({"ok": True, "fecha_display": None})
            messages.info(
                request,
                f"Se ha quitado la fecha para la Evaluación N°{evaluacion.numero_evaluacion}.",
            )

    return redirect("carrera_academica:detalle_ca", pk=evaluacion.carrera_academica.pk)


@login_required
def borrar_archivo_formulario_view(request, pk):
    """
    Borra el archivo de un formulario.
    """
    formulario = get_object_or_404(Formulario, pk=pk)
    ca = formulario.carrera_academica

    if request.method == "POST":
        if formulario.archivo:
            # Guardar nombre para el mensaje
            nombre_archivo = formulario.archivo.name

            # Borrar archivo físico
            formulario.archivo.delete(save=False)

            # Actualizar estado
            formulario.archivo = None
            formulario.estado = 'PEN'
            formulario.fecha_entrega = None
            formulario.save()

            messages.success(
                request,
                f"Archivo eliminado: {formulario.tipo_formulario}"
            )
        else:
            messages.warning(request, "El formulario no tenía archivo adjunto")

        return redirect('carrera_academica:detalle_ca', pk=ca.pk)

    # GET - Confirmación
    context = {
        'formulario': formulario,
        'ca': ca,
    }
    return render(request, 'carrera_academica/confirmar_borrar_archivo.html', context)


@login_required
def gestionar_anios_ca_view(request, pk):
    """
    Vista unificada para:
    - Ver estado de años (pendiente/pausado/evaluado)
    - Pausar/Reactivar años
    - Agregar años nuevos al expediente
    """
    ca = get_object_or_404(CarreraAcademica, pk=pk)

    if request.method == 'POST':
        accion = request.POST.get('accion')

        # ===== PAUSAR AÑO =====
        if accion == 'pausar':
            anio = int(request.POST.get('anio'))
            motivo = request.POST.get('motivo', '').strip()

            if not motivo:
                messages.error(
                    request, "Debe proporcionar un motivo para pausar el año")
                return redirect('carrera_academica:gestionar_anios_ca', pk=pk)

            exito, mensaje = ca.pausar_anio(
                anio, motivo, request.user.username)

            if exito:
                messages.success(request, mensaje)
            else:
                messages.error(request, mensaje)

        # ===== REACTIVAR AÑO =====
        elif accion == 'reactivar':
            anio = int(request.POST.get('anio'))
            exito, mensaje = ca.reactivar_anio(anio)

            if exito:
                messages.success(request, mensaje)
            else:
                messages.error(request, mensaje)

        # ===== AGREGAR AÑOS NUEVOS =====
        elif accion == 'agregar_anios':
            anios_seleccionados = request.POST.getlist('anios')

            if not anios_seleccionados:
                messages.error(request, "Debe seleccionar al menos un año")
                return redirect('carrera_academica:gestionar_anios_ca', pk=pk)

            total_formularios = 0
            anios_agregados = []

            for anio_str in anios_seleccionados:
                anio = int(anio_str)
                creados, mensaje = ca.agregar_formularios_para_anio(anio)
                total_formularios += creados
                anios_agregados.append(anio)

            messages.success(
                request,
                f"Se agregaron {len(anios_agregados)} año(s) con {total_formularios} formularios"
            )

        return redirect('carrera_academica:gestionar_anios_ca', pk=pk)

    # ===== GET: Preparar contexto =====
    resumen = ca.get_resumen_anios()

    # Calcular años disponibles para agregar
    start_year = ca.fecha_inicio.year
    current_year = timezone.now().year
    max_year = current_year + 5  # Permitir agregar hasta 5 años adelante

    # Años que ya tienen formularios o están en el resumen
    anios_existentes = {item['anio'] for item in resumen}

    # Años disponibles = todos - los que ya existen
    anios_disponibles = []
    for anio in range(start_year, max_year + 1):
        if anio not in anios_existentes:
            anios_disponibles.append({
                'anio': anio,
                'es_actual': anio == current_year,
                'es_pasado': anio < current_year,
                'es_futuro': anio > current_year,
            })

    context = {
        'ca': ca,
        'resumen': resumen,
        'anios_disponibles': anios_disponibles,
        'anio_actual': current_year,
        'tiene_anios_disponibles': len(anios_disponibles) > 0,
    }

    return render(request, 'carrera_academica/gestionar_anios.html', context)


@login_required
def gestionar_formularios_anio_view(request, pk, anio):
    """
    Vista para gestionar formularios de un año específico.
    Permite agregar formularios a años nuevos.
    """
    ca = get_object_or_404(CarreraAcademica, pk=pk)

    # Verificar que el año esté en rango
    start_year = ca.fecha_inicio.year
    current_year = timezone.now().year

    if not (start_year <= anio <= current_year):
        messages.error(request, f"El año {anio} está fuera del rango de la CA")
        return redirect('carrera_academica:detalle_ca', pk=pk)

    if request.method == 'POST':
        # Crear formularios faltantes para este año
        tipos_seleccionados = request.POST.getlist('tipos_formularios')

        for tipo in tipos_seleccionados:
            # Verificar que no exista ya
            existe = Formulario.objects.filter(
                carrera_academica=ca,
                tipo_formulario=tipo,
                anio_correspondiente=anio
            ).exists()

            if not existe:
                Formulario.objects.create(
                    carrera_academica=ca,
                    tipo_formulario=tipo,
                    anio_correspondiente=anio
                )

        messages.success(request, f"Formularios agregados para el año {anio}")
        return redirect('carrera_academica:gestionar_formularios_anio', pk=pk, anio=anio)

    # GET - Obtener formularios existentes para este año
    formularios_anio = ca.formularios.filter(
        anio_correspondiente=anio
    ).order_by('tipo_formulario')

    # Tipos de formularios anuales
    tipos_anuales = ['F04', 'F05', 'F06', 'F07', 'F13', 'ENC']

    # Verificar cuáles faltan
    tipos_existentes = set(f.tipo_formulario for f in formularios_anio)
    tipos_faltantes = [t for t in tipos_anuales if t not in tipos_existentes]

    context = {
        'ca': ca,
        'anio': anio,
        'formularios': formularios_anio,
        'tipos_faltantes': tipos_faltantes,
    }

    return render(request, 'carrera_academica/gestionar_formularios_anio.html', context)


@login_required
def archivar_ca_view(request, pk):
    ca = get_object_or_404(CarreraAcademica, pk=pk)
    es_fetch = request.headers.get("X-Requested-With") == "fetch"

    if request.method == "POST":
        motivo = request.POST.get("motivo_archivo")
        observaciones = request.POST.get("observaciones_archivo", "")

        if not motivo:
            if es_fetch:
                return JsonResponse({"ok": False, "error": "Debe seleccionar un motivo de archivo."}, status=400)
            messages.error(request, "Debe seleccionar un motivo de archivo")
            return redirect("carrera_academica:archivar_ca", pk=pk)

        exito, mensaje = CAService.archivar(ca, motivo, observaciones)

        if exito:
            if es_fetch:
                return JsonResponse({
                    "ok": True,
                    "estado_display": ca.get_estado_display(),
                    "motivo_display": ca.get_motivo_archivo_display(),
                })
            messages.success(request, mensaje)
            return redirect("carrera_academica:detalle_ca", pk=pk)
        else:
            if es_fetch:
                return JsonResponse({"ok": False, "error": mensaje}, status=400)
            messages.error(request, mensaje)
            return redirect("carrera_academica:archivar_ca", pk=pk)

    context = {
        "ca": ca,
        "motivo_choices": CarreraAcademica.MOTIVO_ARCHIVO_CHOICES,
    }
    return render(request, "carrera_academica/archivar_ca.html", context)


@login_required
def gestionar_ca_archivada_view(request, pk):
    ca = get_object_or_404(CarreraAcademica, pk=pk)

    if request.method == "POST":
        accion = request.POST.get("accion")

        if accion == "actualizar_motivo":
            nuevo_motivo = request.POST.get("motivo_archivo")
            nuevas_observaciones = request.POST.get(
                "observaciones_archivo", "")

            if not nuevo_motivo:
                messages.error(request, "Debe seleccionar un motivo")
                return redirect("carrera_academica:gestionar_ca_archivada", pk=pk)

            ca.motivo_archivo = nuevo_motivo
            ca.observaciones_archivo = nuevas_observaciones
            ca.save()
            messages.success(
                request, f"Motivo actualizado: {ca.get_motivo_archivo_display()}")

        elif accion == "cerrar_definitivamente":
            resultado = request.POST.get("resultado_cierre")
            observaciones = request.POST.get("observaciones_cierre", "")

            if not resultado:
                messages.error(
                    request, "Debe seleccionar un resultado de cierre")
                return redirect("carrera_academica:gestionar_ca_archivada", pk=pk)

            exito, mensaje = CAService.cerrar_archivada(
                ca, resultado, observaciones, request.user)

            if exito:
                messages.success(
                    request, f"CA cerrada definitivamente: {mensaje}")
            else:
                messages.error(request, mensaje)
                return redirect("carrera_academica:gestionar_ca_archivada", pk=pk)

        return redirect("carrera_academica:detalle_ca", pk=pk)

    context = {
        "ca": ca,
        "motivo_choices": CarreraAcademica.MOTIVO_ARCHIVO_CHOICES,
        "resultado_choices": [
            ("renuncia", "Renuncia Efectivizada"),
            ("jubilacion", "Jubilación Efectivizada"),
            ("no_aprobada", "CA Vencida sin Aprobación"),
        ],
    }
    return render(request, "carrera_academica/gestionar_ca_archivada.html", context)

# ===== VISTAS DE JURADOS =====


@login_required
def lista_jurados_view(request):
    """Lista de jurados disponibles."""
    departamento_filter = request.GET.get('departamento', '')

    jurados = Jurado.objects.filter(activo=True).select_related(
        'profesor_titular_1', 'profesor_titular_2', 'profesor_titular_3',
        'veedor_graduado_titular', 'veedor_estudiante_titular'
    ).prefetch_related('carreras_academicas')

    if departamento_filter:
        jurados = jurados.filter(departamento=departamento_filter)

    jurados = jurados.order_by('-fecha_creacion')

    context = {
        'jurados': jurados,
        'departamento_filter': departamento_filter,
        'departamentos': Asignatura.DEPTO_CHOICES,
    }
    return render(request, 'carrera_academica/jurados/lista.html', context)


@login_required
def crear_jurado_view(request):
    """Crear nuevo jurado."""
    if request.method == 'POST':
        form = JuradoForm(request.POST)
        if form.is_valid():
            jurado = form.save()
            messages.success(
                request, f"Jurado '{jurado.nombre}' creado exitosamente")
            return redirect('carrera_academica:lista_jurados')
    else:
        form = JuradoForm()

    context = {
        'form': form,
        'titulo': 'Crear Jurado',
    }
    return render(request, 'carrera_academica/jurados/form.html', context)


@login_required
def editar_jurado_view(request, pk):
    """Editar jurado existente."""
    jurado = get_object_or_404(Jurado, pk=pk)

    if request.method == 'POST':
        form = JuradoForm(request.POST, instance=jurado)
        if form.is_valid():
            jurado = form.save()
            messages.success(
                request, f"Jurado '{jurado.nombre}' actualizado exitosamente")
            return redirect('carrera_academica:lista_jurados')
    else:
        form = JuradoForm(instance=jurado)

    context = {
        'form': form,
        'jurado': jurado,
        'titulo': f'Editar Jurado: {jurado.nombre}',
    }
    return render(request, 'carrera_academica/jurados/form.html', context)


@login_required
def detalle_jurado_view(request, pk):
    """Detalle de un jurado con las CAs asignadas."""
    jurado = get_object_or_404(
        Jurado.objects.select_related(
            'profesor_titular_1', 'profesor_suplente_1',
            'profesor_titular_2', 'profesor_suplente_2',
            'profesor_titular_3', 'profesor_suplente_3',
            'veedor_graduado_titular', 'veedor_graduado_suplente',
            'veedor_estudiante_titular', 'veedor_estudiante_suplente'
        ).prefetch_related(
            'carreras_academicas__cargo__docente',
            'carreras_academicas__cargo__asignatura'
        ),
        pk=pk
    )

    context = {'jurado': jurado}
    return render(request, 'carrera_academica/jurados/detalle.html', context)


@login_required
def descargar_planilla_jurado_view(request, pk):
    """Descarga la planilla PDF de composición del jurado."""
    jurado = get_object_or_404(
        Jurado.objects.select_related(
            'profesor_titular_1', 'profesor_suplente_1',
            'profesor_titular_2__universidad', 'profesor_suplente_2__universidad',
            'profesor_titular_3__universidad', 'profesor_suplente_3__universidad',
            'veedor_graduado_titular', 'veedor_graduado_suplente',
            'veedor_estudiante_titular', 'veedor_estudiante_suplente',
        ),
        pk=pk
    )

    pdf_file = PDFService.generar_planilla_jurado(jurado)

    if not pdf_file:
        messages.error(request, "No se pudo generar la planilla del jurado.")
        return redirect('carrera_academica:detalle_jurado', pk=pk)

    nombre_archivo = slugify(jurado.nombre) or f"jurado_{pk}"
    response = HttpResponse(pdf_file, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="planilla_jurado_{nombre_archivo}.pdf"'
    return response


@login_required
def asignar_jurado_ca_view(request, ca_pk):
    """Asignar un jurado a una CA."""
    ca = get_object_or_404(CarreraAcademica, pk=ca_pk)

    if request.method == 'POST':
        jurado_id = request.POST.get('jurado_id')

        if jurado_id:
            jurado = get_object_or_404(Jurado, pk=jurado_id, activo=True)
            ca.jurado = jurado
            ca.save()
            messages.success(
                request, f"Jurado '{jurado.nombre}' asignado correctamente")
        else:
            ca.jurado = None
            ca.save()
            messages.info(request, "Jurado desvinculado de la CA")

        return redirect('carrera_academica:detalle_ca', pk=ca.pk)

    # Filtrar jurados por el departamento de la asignatura
    departamento = ca.cargo.asignatura.departamento
    jurados = Jurado.objects.filter(
        activo=True,
        departamento=departamento
    ).order_by('nombre')

    context = {
        'ca': ca,
        'jurados': jurados,
    }
    return render(request, 'carrera_academica/jurados/asignar.html', context)


# ===== VISTAS DE VEEDORES =====

@login_required
def lista_veedores_graduados_view(request):
    """Lista de veedores graduados."""
    veedores = VeedorGraduado.objects.all().order_by('apellido', 'nombre')
    return render(request, 'carrera_academica/veedores/lista_graduados.html', {'veedores': veedores})


@login_required
def crear_veedor_graduado_view(request):
    """Crear veedor graduado."""
    if request.method == 'POST':
        form = VeedorGraduadoForm(request.POST)
        if form.is_valid():
            veedor = form.save()
            messages.success(
                request, f"Veedor graduado '{veedor}' creado exitosamente")
            return redirect('carrera_academica:lista_veedores_graduados')
    else:
        form = VeedorGraduadoForm()

    return render(request, 'carrera_academica/veedores/form_graduado.html', {'form': form})


@login_required
def lista_veedores_estudiantes_view(request):
    """Lista de veedores estudiantes."""
    veedores = VeedorEstudiante.objects.all().order_by('apellido', 'nombre')
    return render(request, 'carrera_academica/veedores/lista_estudiantes.html', {'veedores': veedores})


@login_required
def crear_veedor_estudiante_view(request):
    """Crear veedor estudiante."""
    if request.method == 'POST':
        form = VeedorEstudianteForm(request.POST)
        if form.is_valid():
            veedor = form.save()
            messages.success(
                request, f"Veedor estudiante '{veedor}' creado exitosamente")
            return redirect('carrera_academica:lista_veedores_estudiantes')
    else:
        form = VeedorEstudianteForm()

    return render(request, 'carrera_academica/veedores/form_estudiante.html', {'form': form})


@login_required
def crear_veedor_graduado_ajax(request):
    """Crear veedor graduado vía AJAX."""
    if request.method == 'POST':
        form = VeedorGraduadoForm(request.POST)
        if form.is_valid():
            veedor = form.save()
            return JsonResponse({
                'success': True,
                'veedor': {
                    'id': veedor.pk,
                    'texto': str(veedor)
                }
            })
        else:
            return JsonResponse({
                'success': False,
                'error': str(form.errors)
            })
    return JsonResponse({'success': False, 'error': 'Método no permitido'})


@login_required
def crear_veedor_estudiante_ajax(request):
    """Crear veedor estudiante vía AJAX."""
    if request.method == 'POST':
        form = VeedorEstudianteForm(request.POST)
        if form.is_valid():
            veedor = form.save()
            return JsonResponse({
                'success': True,
                'veedor': {
                    'id': veedor.pk,
                    'texto': str(veedor)
                }
            })
        else:
            return JsonResponse({
                'success': False,
                'error': str(form.errors)
            })
    return JsonResponse({'success': False, 'error': 'Método no permitido'})


# ===== VISTAS JURADOS EXTERNOS =====

@login_required
def lista_jurados_externos_view(request):
    """Lista de jurados externos."""
    universidad_filter = request.GET.get('universidad', '')

    jurados = JuradoExterno.objects.filter(
        activo=True).select_related('universidad')

    if universidad_filter:
        jurados = jurados.filter(universidad__sigla=universidad_filter)

    universidades = Universidad.objects.all().order_by('sigla')

    context = {
        'jurados': jurados,
        'universidades': universidades,
        'universidad_filter': universidad_filter,
    }
    return render(request, 'carrera_academica/jurados_externos/lista.html', context)


@login_required
def crear_jurado_externo_view(request):
    """Crear jurado externo."""
    if request.method == 'POST':
        form = JuradoExternoForm(request.POST, request.FILES)
        if form.is_valid():
            jurado = form.save()
            messages.success(
                request, f'Jurado externo {jurado} creado exitosamente.')
            return redirect('carrera_academica:lista_jurados_externos')
    else:
        form = JuradoExternoForm()

    context = {
        'form': form,
        'titulo': 'Crear Jurado Externo'
    }
    return render(request, 'carrera_academica/jurados_externos/form.html', context)


@login_required
def editar_jurado_externo_view(request, pk):
    """Editar jurado externo."""
    jurado = get_object_or_404(JuradoExterno, pk=pk)

    if request.method == 'POST':
        form = JuradoExternoForm(request.POST, request.FILES, instance=jurado)
        if form.is_valid():
            form.save()
            messages.success(request, f'Jurado externo {jurado} actualizado.')
            return redirect('carrera_academica:lista_jurados_externos')
    else:
        form = JuradoExternoForm(instance=jurado)

    context = {
        'form': form,
        'jurado': jurado,
        'titulo': f'Editar: {jurado}'
    }
    return render(request, 'carrera_academica/jurados_externos/form.html', context)


@login_required
def detalle_jurado_externo_view(request, pk):
    """Detalle de jurado externo."""
    jurado = get_object_or_404(JuradoExterno, pk=pk)

    # Jurados donde participa como titular 2
    jurados_titular_2 = Jurado.objects.filter(profesor_titular_2=jurado)

    # Jurados donde participa como titular 3
    jurados_titular_3 = Jurado.objects.filter(profesor_titular_3=jurado)

    # Jurados como suplente
    jurados_suplente_2 = Jurado.objects.filter(profesor_suplente_2=jurado)
    jurados_suplente_3 = Jurado.objects.filter(profesor_suplente_3=jurado)

    context = {
        'jurado': jurado,
        'jurados_titular_2': jurados_titular_2,
        'jurados_titular_3': jurados_titular_3,
        'jurados_suplente_2': jurados_suplente_2,
        'jurados_suplente_3': jurados_suplente_3,
    }
    return render(request, 'carrera_academica/jurados_externos/detalle.html', context)


@login_required
def crear_jurado_externo_ajax(request):
    """Crear jurado externo vía AJAX desde modal."""
    if request.method == 'POST':
        form = JuradoExternoForm(request.POST, request.FILES)
        if form.is_valid():
            jurado = form.save()
            return JsonResponse({
                'success': True,
                'jurado': {
                    'id': jurado.pk,
                    'texto': str(jurado)
                }
            })
        else:
            return JsonResponse({
                'success': False,
                'error': form.errors
            })
    return JsonResponse({'success': False, 'error': 'Método no permitido'})


# ===== VISTAS UNIVERSIDADES =====

@login_required
def lista_universidades_view(request):
    """Lista de universidades."""
    universidades = Universidad.objects.all().order_by('sigla')

    context = {
        'universidades': universidades,
    }
    return render(request, 'carrera_academica/universidades/lista.html', context)


@login_required
def crear_universidad_view(request):
    """Crear universidad."""
    if request.method == 'POST':
        form = UniversidadForm(request.POST)
        if form.is_valid():
            universidad = form.save()
            messages.success(
                request, f'Universidad {universidad} creada exitosamente.')
            return redirect('carrera_academica:lista_universidades')
    else:
        form = UniversidadForm()

    context = {
        'form': form,
        'titulo': 'Crear Universidad'
    }
    return render(request, 'carrera_academica/universidades/form.html', context)


@login_required
def editar_universidad_view(request, pk):
    """Editar universidad."""
    universidad = get_object_or_404(Universidad, pk=pk)

    if request.method == 'POST':
        form = UniversidadForm(request.POST, instance=universidad)
        if form.is_valid():
            form.save()
            messages.success(
                request, f'Universidad {universidad} actualizada.')
            return redirect('carrera_academica:lista_universidades')
    else:
        form = UniversidadForm(instance=universidad)

    context = {
        'form': form,
        'universidad': universidad,
        'titulo': f'Editar: {universidad}'
    }
    return render(request, 'carrera_academica/universidades/form.html', context)
