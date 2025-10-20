# equivalencias/views.py
# Django imports
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.conf import settings
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.core.mail import EmailMessage
from django.db.models import Case, When, Value
from django.contrib import messages
from django.db.models import Count, Avg, F, DurationField
from django.db.models.functions import TruncMonth, Trim, ExtractMonth
from django.utils import timezone

# Library imports
import io
import mimetypes
import os
from docx import Document
from weasyprint import HTML
from datetime import date

# Model imports
from .models import (
    Estudiante,
    AsignaturaParaEquivalencia,
    SolicitudEquivalencia,
    DetalleSolicitud,
    DocumentoAdjunto,
)


def _enviar_email_catedra(detalle_solicitud):
    """
    Función auxiliar que genera el Word, los adjuntos y envía el email
    para un detalle de solicitud específico.
    """
    solicitud = detalle_solicitud.id_solicitud
    asignatura = detalle_solicitud.id_asignatura
    estudiante = solicitud.id_estudiante

    # Obtener el email del responsable
    responsable = asignatura.docente_responsable
    correo_principal = responsable.correos.filter(principal=True).first()

    if not correo_principal:
        raise ValueError(f"No se encontró correo principal para {responsable}")

    # --- 2. OBTÉN Y FORMATEA LA FECHA ACTUAL ---
    hoy = date.today()
    meses = (
        "Enero",
        "Febrero",
        "Marzo",
        "Abril",
        "Mayo",
        "Junio",
        "Julio",
        "Agosto",
        "Septiembre",
        "Octubre",
        "Noviembre",
        "Diciembre",
    )
    fecha_actual_texto = f"{hoy.day} de {meses[hoy.month - 1]} de {hoy.year}"

    # Personalizar el documento de Word en memoria
    doc = Document("templates_word/planilla_evaluacion.docx")
    for p in doc.paragraphs:
        if "[fecha]" in p.text:
            p.text = p.text.replace("[fecha]", fecha_actual_texto)
        if "[alumno]" in p.text:
            p.text = p.text.replace("[alumno]", estudiante.nombre_completo)
        if "[asignatura]" in p.text:
            p.text = p.text.replace("[asignatura]", asignatura.asignatura.nombre)

    # ... (Si tienes más reemplazos en tablas, también irían aquí)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "[fecha]" in p.text:
                        p.text = p.text.replace("[fecha]", fecha_actual_texto)
                    if "[asignatura]" in p.text:
                        p.text = p.text.replace(
                            "[asignatura]", asignatura.asignatura.nombre
                        )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Preparar y enviar el correo
    email = EmailMessage(
        # Cambiamos el asunto para que sirva en ambos casos
        subject=f"Solicitud de Equivalencia de {estudiante.nombre_completo}",
        # Cambiamos el cuerpo del correo
        body=f"""
        <p>Hola profe, le envío la documentación para dar, si corresponde, la equivalencia de <strong>{asignatura.asignatura.nombre}</strong>, al futuro estudiante <strong>{estudiante.nombre_completo}</strong>.</p>

        <p>Agradeceré que responda a este mismo correo con su dictamen. No es necesario reenviar el archivo.</p>
        
        <p>Sin otro particular.</p>
        <p><br></p>

        <p style='font-size:15px;font-family:"Calibri",sans-serif; color:#174E86;'>
            <strong>Ing. Jorge RONCONI</strong><br>
            Secretario de Departamento
        </p>
        <p style='font-size:15px;font-family:"Calibri",sans-serif;color:#174E86;'>
            <strong>
                Departamento Ingeniería Civil<br>
                Universidad Tecnológica Nacional<br>
                Facultad Regional La Plata
            </strong>
        </p>
        <p style='font-size:15px;font-family:"Calibri",sans-serif;text-align:center;color:#70AD47;'>
            Universidad Publica, Gratuita y de Calidad.
        </p>
        """,
        from_email=None,
        to=[correo_principal.email],
    )

    for documento_adjunto in solicitud.documentoadjunto_set.all():
        content_type = (
            mimetypes.guess_type(documento_adjunto.archivo.name)[0]
            or "application/octet-stream"
        )
        email.attach(
            documento_adjunto.archivo.name.split("/")[-1],
            documento_adjunto.archivo.read(),
            content_type,
        )

    email.attach(
        f"Planilla_{estudiante.dni_pasaporte}_{asignatura.asignatura.nombre}.docx",
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    email.content_subtype = "html"  # Para enviar el cuerpo como HTML
    email.send()


def _calculate_statistics(solicitudes_qs, is_historical_view=False):
    """
    Recibe un QuerySet de solicitudes y calcula todas las métricas.
    Devuelve un diccionario con todos los resultados listos para la plantilla.
    """
    solicitudes_ids = solicitudes_qs.values_list("id", flat=True)
    detalles_qs = DetalleSolicitud.objects.filter(id_solicitud__in=solicitudes_ids)

    # --- Métricas de Volumen ---
    total_solicitudes = solicitudes_qs.count()
    total_asignaturas_procesadas = detalles_qs.count()

    # --- Gráfico 1: Solicitudes por Período ---

    meses_labels_es = [
        "Enero",
        "Febrero",
        "Marzo",
        "Abril",
        "Mayo",
        "Junio",
        "Julio",
        "Agosto",
        "Septiembre",
        "Octubre",
        "Noviembre",
        "Diciembre",
    ]

    if is_historical_view:
        # --- NUEVA LÓGICA PARA EL PROMEDIO HISTÓRICO ---
        num_years = solicitudes_qs.dates("fecha_inicio", "year").count()
        num_years = num_years if num_years > 0 else 1  # Evitar división por cero

        # Solicitudes nuevas
        nuevas_por_mes_total = (
            solicitudes_qs.annotate(month=ExtractMonth("fecha_inicio"))
            .values("month")
            .annotate(total=Count("id"))
            .order_by("month")
        )
        # Completadas
        completadas_por_mes_total = (
            solicitudes_qs.filter(
                estado_general="Completada", fecha_completada__isnull=False
            )
            .annotate(month=ExtractMonth("fecha_completada"))
            .values("month")
            .annotate(total=Count("id"))
            .order_by("month")
        )

        # Preparamos los datos promediados
        solicitudes_data_promedio = [0] * 12
        nuevas_data_promedio = [0] * 12
        completadas_data_promedio = [0] * 12

        for item in nuevas_por_mes_total:
            solicitudes_data_promedio[item["month"] - 1] = item["total"] / num_years
            nuevas_data_promedio[item["month"] - 1] = item["total"] / num_years

        for item in completadas_por_mes_total:
            completadas_data_promedio[item["month"] - 1] = item["total"] / num_years

        meses_labels = meses_labels_es
        solicitudes_data = solicitudes_data_promedio
        balance_labels = meses_labels_es
        nuevas_data = nuevas_data_promedio
        completadas_data = completadas_data_promedio

    else:
        solicitudes_por_mes = (
            solicitudes_qs.annotate(mes=TruncMonth("fecha_inicio"))
            .values("mes")
            .annotate(total=Count("id"))
            .order_by("mes")
        )
        meses_labels = [
            s["mes"].strftime("%B %Y").capitalize() for s in solicitudes_por_mes
        ]
        solicitudes_data = [s["total"] for s in solicitudes_por_mes]

        # --- Gráfico 2: Balance Mensual ---
        completadas_por_mes = (
            solicitudes_qs.filter(
                estado_general="Completada", fecha_completada__isnull=False
            )
            .annotate(mes=TruncMonth("fecha_completada"))
            .values("mes")
            .annotate(total=Count("id"))
            .order_by("mes")
        )
        balance_dict = {
            s["mes"]: {"nuevas": s["total"], "completadas": 0}
            for s in solicitudes_por_mes
        }
        for item in completadas_por_mes:
            if item["mes"] in balance_dict:
                balance_dict[item["mes"]]["completadas"] = item["total"]
        balance_labels = [
            mes.strftime("%B %Y").capitalize() for mes in sorted(balance_dict.keys())
        ]
        nuevas_data = [
            balance_dict[mes]["nuevas"] for mes in sorted(balance_dict.keys())
        ]
        completadas_data = [
            balance_dict[mes]["completadas"] for mes in sorted(balance_dict.keys())
        ]

    # --- Métricas de Resultados ---
    dictamenes_data = (
        detalles_qs.annotate(estado_limpio=Trim("estado_asignatura"))
        .exclude(estado_limpio__in=["Enviada a Cátedra", "Pendiente de envío"])
        .values("estado_limpio")
        .annotate(total=Count("id"))
        .order_by("total")
    )
    dictamen_labels = [d["estado_limpio"] for d in dictamenes_data]
    dictamen_valores = [d["total"] for d in dictamenes_data]

    # <<< CORRECCIÓN: Usamos .annotate() para crear un alias 'nombre_materia' >>>
    asignaturas_mas_solicitadas = (
        detalles_qs.values("id_asignatura__asignatura__nombre")
        .annotate(
            nombre_materia=F("id_asignatura__asignatura__nombre"), total=Count("id")
        )
        .order_by("-total")
        .values("nombre_materia", "total")[:10]
    )

    asignaturas_problematicas = (
        detalles_qs.filter(estado_asignatura__in=["Denegada", "Requiere PC"])
        .values("id_asignatura__asignatura__nombre", "estado_asignatura")
        .annotate(
            nombre_materia=F("id_asignatura__asignatura__nombre"), total=Count("id")
        )
        .order_by("-total")
        .values("nombre_materia", "estado_asignatura", "total")[:10]
    )

    # --- Métricas de Tiempos (CORREGIDAS) ---
    avg_resolucion = solicitudes_qs.filter(
        estado_general="Completada", fecha_completada__isnull=False
    ).aggregate(avg_diff=Avg(F("fecha_completada") - F("fecha_inicio")))["avg_diff"]

    avg_dictamen = detalles_qs.filter(fecha_dictamen__isnull=False).aggregate(
        avg_diff=Avg(F("fecha_dictamen") - F("id_solicitud__fecha_inicio"))
    )["avg_diff"]

    top_demoras_query = (
        detalles_qs.filter(fecha_dictamen__isnull=False)
        .annotate(duracion=F("fecha_dictamen") - F("id_solicitud__fecha_inicio"))
        .values("id_asignatura__asignatura__nombre")
        .annotate(
            nombre_materia=F("id_asignatura__asignatura__nombre"),
            demora_promedio=Avg("duracion"),
        )
        .order_by("-demora_promedio")
        .values("nombre_materia", "demora_promedio")[:5]
    )
    top_demoras = [
        {"nombre": item["nombre_materia"], "dias": item["demora_promedio"].days}
        for item in top_demoras_query
    ]

    # --- Devolvemos todo en un solo diccionario ---
    return {
        "total_solicitudes": total_solicitudes,
        "total_asignaturas_procesadas": total_asignaturas_procesadas,
        "meses_labels": meses_labels,
        "solicitudes_data": solicitudes_data,
        "balance_labels": balance_labels,
        "nuevas_data": nuevas_data,
        "completadas_data": completadas_data,
        "dictamen_labels": dictamen_labels,
        "dictamen_valores": dictamen_valores,
        "asignaturas_mas_solicitadas": asignaturas_mas_solicitadas,
        "asignaturas_problematicas": asignaturas_problematicas,
        "avg_resolucion_dias": avg_resolucion.days if avg_resolucion else 0,
        "avg_dictamen_dias": avg_dictamen.days if avg_dictamen else 0,
        "top_demoras": top_demoras,
    }


@login_required
def dashboard_view(request):
    # 1. Lógica de Búsqueda
    # Obtiene el término de búsqueda de la URL
    search_query = request.GET.get("q", "")

    solicitudes = SolicitudEquivalencia.objects.all()

    if search_query:
        # Filtra por nombre completo del estudiante (insensible a mayúsculas/minúsculas)
        solicitudes = solicitudes.filter(
            id_estudiante__nombre_completo__icontains=search_query
        )

    # 2. Lógica de Ordenamiento
    # Anotamos cada solicitud con un '1' si está completada, y '0' si no.
    solicitudes = solicitudes.annotate(
        estado_ordenado=Case(
            When(estado_general="Completada", then=Value(1)), default=Value(0)
        )
        # Ordena primero por estado, luego por fecha
    ).order_by("estado_ordenado", "fecha_inicio")

    # 3. Pasamos los datos a la plantilla
    contexto = {
        "solicitudes": solicitudes,
        "search_query": search_query,  # Para mantener el texto en la barra de búsqueda
    }
    return render(request, "equivalencias/dashboard.html", contexto)


@login_required
def solicitud_detalle_view(request, pk):
    solicitud = get_object_or_404(SolicitudEquivalencia, pk=pk)

    # --- LÓGICA DE ACTUALIZACIÓN (POST) CORREGIDA ---
    if request.method == "POST":
        detalle_id = request.POST.get("detalle_id")
        nuevo_estado = request.POST.get("estado_asignatura")

        if detalle_id and nuevo_estado:
            detalle_a_actualizar = get_object_or_404(DetalleSolicitud, pk=detalle_id)
            detalle_a_actualizar.estado_asignatura = nuevo_estado

            # --- LÓGICA REINCORPORADA PARA GUARDAR EL PC ---
            if nuevo_estado == "Requiere PC":
                # Obtenemos el texto del formulario y lo guardamos
                temas_pc = request.POST.get("detalle_pc", "")
                detalle_a_actualizar.detalle_pc = temas_pc
            else:
                # Buena práctica: si el estado no es PC, limpiamos el campo
                detalle_a_actualizar.detalle_pc = None

            estados_finales = ["Aprobada", "Denegada", "Requiere PC"]
            if nuevo_estado in estados_finales:
                detalle_a_actualizar.fecha_dictamen = timezone.now()
            else:
                detalle_a_actualizar.fecha_dictamen = None

            detalle_a_actualizar.save()

            nombre_materia = (
                detalle_a_actualizar.id_asignatura.asignatura.nombre.title()
            )
            messages.success(
                request, f"El estado de '{nombre_materia}' fue actualizado."
            )

        return redirect("solicitud_detalle", pk=pk)

    # --- LÓGICA PARA MOSTRAR LA PÁGINA (GET) ---
    detalles = solicitud.detallesolicitud_set.select_related(
        "id_asignatura__asignatura", "id_asignatura__docente_responsable"
    ).all()

    estado_choices = DetalleSolicitud.ESTADO_ASIGNATURA_CHOICES

    # Lógica para determinar si la solicitud está completa
    solicitud_completa = False
    if detalles.exists():
        estados_finales = ["Aprobada", "Denegada", "Requiere PC"]
        solicitud_completa = all(
            d.estado_asignatura in estados_finales for d in detalles
        )

    contexto = {
        "solicitud": solicitud,
        "detalles": detalles,
        "estado_choices": estado_choices,
        "solicitud_completa": solicitud_completa,
    }
    return render(request, "equivalencias/solicitud_detalle.html", contexto)


@login_required
def crear_solicitud_view(request):
    if request.method == "POST":
        # --- LÓGICA PARA OBTENER O CREAR EL ESTUDIANTE ---
        estudiante_id = request.POST.get("estudiante")
        nuevo_estudiante_nombre = request.POST.get("nuevo_estudiante_nombre")
        nuevo_estudiante_dni = request.POST.get("nuevo_estudiante_dni")

        estudiante = None
        if estudiante_id:
            estudiante = get_object_or_404(Estudiante, pk=estudiante_id)
        elif nuevo_estudiante_nombre and nuevo_estudiante_dni:
            estudiante, created = Estudiante.objects.get_or_create(
                dni_pasaporte=nuevo_estudiante_dni,
                defaults={"nombre_completo": nuevo_estudiante_nombre},
            )

        if not estudiante:
            messages.error(
                request, "Debe seleccionar un estudiante existente o crear uno nuevo."
            )
            return redirect("crear_solicitud")

        # --- CREACIÓN DE LA SOLICITUD Y DOCUMENTOS ---
        asignatura_ids = request.POST.getlist("asignaturas")
        documentos = request.FILES.getlist("documentacion")

        nueva_solicitud = SolicitudEquivalencia.objects.create(id_estudiante=estudiante)
        for doc in documentos:
            DocumentoAdjunto.objects.create(solicitud=nueva_solicitud, archivo=doc)

        # --- Bucle para procesar cada asignatura y enviar correo ---
        for asig_id in asignatura_ids:
            asig_para_equiv = get_object_or_404(AsignaturaParaEquivalencia, pk=asig_id)

            # 1. Creamos el registro del detalle de la solicitud
            detalle_solicitud = DetalleSolicitud.objects.create(
                id_solicitud=nueva_solicitud, id_asignatura=asig_para_equiv
            )

            # 2. Enviamos el correo usando la función auxiliar
            try:
                # La función se encarga de todo: generar Word, adjuntar y enviar.
                _enviar_email_catedra(detalle_solicitud)
            except Exception as e:
                # Si algo falla en el envío, informamos al usuario.
                messages.error(
                    request,
                    f"No se pudo enviar el correo para {asig_para_equiv.asignatura.nombre}. Error: {e}",
                )

        messages.success(
            request, "Solicitud creada y notificaciones enviadas exitosamente."
        )
        return redirect("dashboard")

    # --- Lógica para mostrar el formulario (GET) ---
    estudiantes = Estudiante.objects.all()
    asignaturas = (
        AsignaturaParaEquivalencia.objects.all()
        .select_related("asignatura")
        .order_by("asignatura__nombre")
    )

    contexto = {
        "estudiantes": estudiantes,
        "asignaturas": asignaturas,
    }
    return render(request, "equivalencias/crear_solicitud.html", contexto)


@login_required
def generar_acta_pdf_view(request, pk):
    """
    Genera un acta en PDF para una solicitud de equivalencia,
    incluyendo una imagen de firma.
    """
    solicitud = get_object_or_404(SolicitudEquivalencia, pk=pk)
    detalles = solicitud.detallesolicitud_set.all()

    # 1. Construir la ruta absoluta al archivo de la firma
    # (Asegúrate de que la imagen 'firma_holografica.png' exista en 'static/images/')
    try:
        image_path_os = os.path.join(
            settings.STATICFILES_DIRS[0], "images", "firma_holografica.png"
        )
    except (IndexError, AttributeError):
        # Manejo de error si STATICFILES_DIRS no está configurado
        image_path_os = None  # O una ruta a una imagen por defecto

    image_path = image_path_os.replace("\\", "/")

    # 2. Preparar el contexto completo para la plantilla
    context = {
        "solicitud": solicitud,
        "detalles": detalles,
        # Añadimos la ruta de la imagen al contexto. El prefijo 'file://' es clave.
        #'signature_image_path': f'file://{image_path}' if image_path and os.path.exists(image_path) else None,
        "signature_image_path": "/static/images/firma_holografica.png",  # Ruta relativa para WeasyPrint
    }
    print(f"DEBUG: image_path = {context['signature_image_path']}"),

    # 3. Renderizar la plantilla HTML a un string con el contexto actualizado
    html_string = render_to_string("equivalencias/acta_template.html", context)
    # Le decimos a WeasyPrint cuál es la dirección de nuestro sitio.
    # request.build_absolute_uri('/') obtiene algo como 'http://127.0.0.1:8000'
    base_url = request.build_absolute_uri("/")

    # 4. Generar el PDF
    # La librería WeasyPrint usará la ruta absoluta para encontrar y embeber la imagen
    pdf_file = HTML(string=html_string, base_url=base_url).write_pdf()

    # 5. Devolver el PDF como una respuesta para descargar
    response = HttpResponse(pdf_file, content_type="application/pdf")
    response["Content-Disposition"] = (
        f'attachment; filename="acta_{solicitud.id_estudiante.dni_pasaporte}.pdf"'
    )
    return response


@login_required
def finalizar_solicitud_view(request, pk):
    solicitud = get_object_or_404(SolicitudEquivalencia, pk=pk)
    if request.method == "POST":
        acta_firmada_file = request.FILES.get("acta_firmada")
        if acta_firmada_file:
            solicitud.acta_firmada = acta_firmada_file
            solicitud.save()

            # Enviar correo al Dpto. de Alumnos
            email = EmailMessage(
                subject=f"Resolución de Equivalencias - {solicitud.id_estudiante.nombre_completo}",
                body="Se adjunta el acta final de equivalencias para su registro en el legajo del estudiante.",
                from_email=None,
                to=["alumnos@frlp.utn.edu.ar"],  # <-- CAMBIA ESTE EMAIL
            )
            email.attach_file(solicitud.acta_firmada.path)
            email.send()

            # Cambiar estado y archivar
            solicitud.estado_general = "Completada"
            solicitud.fecha_completada = timezone.now()
            solicitud.save()

            # <-- Mensaje mejorado
            messages.success(
                request, "Solicitud finalizada, notificada y archivada correctamente."
            )
            return redirect("dashboard")

    return redirect("solicitud_detalle", pk=pk)  # Redirigir si no es POST


@login_required
def reenviar_email_asignatura_view(request, pk, detalle_pk):
    detalle = get_object_or_404(DetalleSolicitud, pk=detalle_pk)

    try:
        _enviar_email_catedra(detalle)
        messages.success(
            request,
            f"Correo reenviado exitosamente a {detalle.id_asignatura.email_responsable}.",
        )
    except Exception as e:
        messages.error(request, f"Error al reenviar el correo: {e}")

    return redirect("solicitud_detalle", pk=pk)


@login_required
def reenviar_pendientes_view(request, pk):
    solicitud = get_object_or_404(SolicitudEquivalencia, pk=pk)

    # Filtramos solo las asignaturas que aún no tienen respuesta
    pendientes = solicitud.detallesolicitud_set.filter(
        estado_asignatura="Enviada a Cátedra"
    )

    if not pendientes:
        messages.warning(
            request, "No hay correos pendientes de respuesta para reenviar."
        )
        return redirect("solicitud_detalle", pk=pk)

    contador = 0
    for detalle in pendientes:
        try:
            _enviar_email_catedra(detalle)
            contador += 1
        except Exception as e:
            messages.error(
                request,
                f"Falló el envío a {detalle.id_asignatura.docente_responsable}: {e}",
            )
            continue  # Continúa con el siguiente aunque uno falle

    messages.success(
        request, f"Se reenviaron {contador} correos a las cátedras pendientes."
    )
    return redirect("solicitud_detalle", pk=pk)


@login_required
def estadisticas_view(request):
    selected_year = request.GET.get("year")
    available_years_dates = SolicitudEquivalencia.objects.dates(
        "fecha_inicio", "year", order="DESC"
    )
    available_years = [d.year for d in available_years_dates]

    solicitudes_base = SolicitudEquivalencia.objects.all()
    # Título por defecto
    titulo_periodo = "Promedio Histórico (todos los años)"
    is_historical = True  # Flag para la función de cálculo

    if selected_year and selected_year.isdigit():
        if int(selected_year) in available_years:
            solicitudes_base = solicitudes_base.filter(
                fecha_inicio__year=int(selected_year)
            )
            titulo_periodo = f"Año {selected_year}"
            is_historical = False  # Cambiamos el flag

    # Pasamos el flag a la función de cálculo
    stats = _calculate_statistics(solicitudes_base, is_historical_view=is_historical)

    contexto = {
        "stats": stats,
        "available_years": available_years,
        "selected_year": int(selected_year) if selected_year else None,
        "titulo_periodo": titulo_periodo,
    }

    return render(request, "equivalencias/estadisticas.html", contexto)
